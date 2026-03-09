import docx
import os

def extract_comprehensive_text(file_path):
    doc = docx.Document(file_path)
    content = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))
                
    return '\n'.join(content)

base_path = r'C:\Users\rguti\Petral.MARK\Infomemo.Petral\Documentos Research\Sandwell'
file1 = os.path.join(base_path, 'scan2_merged-pages-1.docx')
file2 = os.path.join(base_path, 'scan2_merged-pages-2.docx')

print(f"Extracting from {file1}...")
text1 = extract_comprehensive_text(file1)
print(f"Extracting from {file2}...")
text2 = extract_comprehensive_text(file2)

output_file = os.path.join(base_path, 'extracted_text_v2_comprehensive.md')
with open(output_file, 'w', encoding='utf-8') as f:
    f.write("# COMPREHENSIVE EXTRACTION - Sandwell\n\n")
    f.write("## FILE 1\n\n")
    f.write(text1)
    f.write("\n\n## FILE 2\n\n")
    f.write(text2)

print(f"Comprehensive extraction saved to: {output_file}")
print(f"Total characters: {len(text1) + len(text2)}")
