import docx
from docx.shared import Pt, RGBColor
import os
import sys

def analyze_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Error: File not found {file_path}")
        return

    doc = docx.Document(file_path)
    
    print(f"--- Analysis of {os.path.basename(file_path)} ---\n")
    
    # Page Margins (approximate from first section)
    for section in doc.sections:
        print("Page Setup:")
        print(f"  Page Width: {section.page_width.inches} inches")
        print(f"  Page Height: {section.page_height.inches} inches")
        print(f"  Left Margin: {section.left_margin.inches} inches")
        print(f"  Right Margin: {section.right_margin.inches} inches")
        print(f"  Top Margin: {section.top_margin.inches} inches")
        print(f"  Bottom Margin: {section.bottom_margin.inches} inches")
        break # Just first section
    
    print("\n--- Styles Used ---")
    styles_used = set()
    for para in doc.paragraphs:
        styles_used.add(para.style.name)
    
    for style in sorted(list(styles_used)):
        print(f"  - {style}")
        
    print("\n--- Content Preview (First 20 paragraphs) ---")
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"[{para.style.name}] {para.text[:100]}...")

    print("\n--- Font & Formatting Checks (Sample) ---")
    # Check a few styles for font details if possible
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Title']:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            font = style.font
            print(f"Style '{style_name}':")
            print(f"  Font: {font.name}")
            print(f"  Size: {font.size.pt if font.size else 'Default'}")
            print(f"  Color: {font.color.rgb if font.color else 'Default'}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        analyze_docx(sys.argv[1])
    else:
        print("Usage: python extract_docx_info.py <docx_path>")
